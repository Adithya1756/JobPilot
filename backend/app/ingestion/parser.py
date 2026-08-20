"""
Document parser - extracts text from PDF and DOCX files.
This is a foundational piece of the ingestion pipeline.
"""

import tempfile
from pathlib import Path
from typing import Optional, Union
import pdfplumber
from docx import Document


class DocumentParser:
    """
    Extracts text content from uploaded documents.

    Supports:
    - PDF files (via pdfplumber)
    - DOCX files (via python-docx)
    - Plain text files

    Why pdfplumber over PyPDF2:
    - Better table extraction (resumes often have tabular layouts)
    - Preserves reading order better
    - Handles multi-column layouts more reliably
    """

    @staticmethod
    def parse_pdf(file_path: Union[str, Path]) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text content
        """
        text_parts = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text from page
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Also extract tables (common in resumes for skills, etc.)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        # Join non-None cell values
                        row_text = " | ".join(cell for cell in row if cell)
                        if row_text.strip():
                            text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    def parse_docx(file_path: Union[str, Path]) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content
        """
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    def parse_text(file_path: Union[str, Path]) -> str:
        """Read a plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @classmethod
    def parse(cls, file_path: Union[str, Path], content_type: Optional[str] = None) -> str:
        """
        Parse a document based on file extension or content type.

        Args:
            file_path: Path to the file
            content_type: Optional MIME type hint

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf" or content_type == "application/pdf":
            return cls.parse_pdf(path)
        elif suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return cls.parse_docx(path)
        elif suffix in (".txt", ".md") or content_type in ("text/plain", "text/markdown"):
            return cls.parse_text(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")


def parse_uploaded_file(file_content: bytes, filename: str) -> str:
    """
    Parse an uploaded file given its bytes and filename.

    This is the main entry point for the ingestion pipeline.
    Creates a temporary file to work with the document parsers.

    Args:
        file_content: Raw bytes of the uploaded file
        filename: Original filename (used to determine type)

    Returns:
        Extracted text content
    """
    suffix = Path(filename).suffix.lower()

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        return DocumentParser.parse(tmp_path)
    finally:
        # Clean up temp file
        Path(tmp_path).unlink(missing_ok=True)
